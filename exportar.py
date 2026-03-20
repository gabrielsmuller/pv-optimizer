import sys, os, re
from docx.shared import Inches
from docx.shared import Pt, RGBColor
from datetime import datetime
from docx import Document
from tkinter import filedialog
import io
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from popups import CustomInfoPopupWithOpen, CustomInfoPopup
import matplotlib.pyplot as plt
import numpy as np

if getattr(sys, 'frozen', False):
    application_path = sys._MEIPASS
else:
    application_path = os.path.dirname(os.path.abspath(__file__))

icon_path = "dimicon.ico"
icon_full_path = os.path.join(application_path, icon_path)

def exportar(self, tipo_sistema, tipo_rede, eficiencia, autonomia, potencia, consumo, painel, dimensionamento, tabela, endereco, coordenadas, orientacao, inclinacao, temperatura, seguranca, tabela_consumo, hsp, modo_calculo, organizacao):

    document = Document()

    # Cabeçalho
    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    path_da_imagem = os.path.join(application_path, "energygroup.png")
    run = header_paragraph.add_run()
    run.add_picture(path_da_imagem, width=Inches(2.0))

    # Adicionar título e data
    document.add_paragraph(f"")
    document.add_heading(f'Simulação de Sistema {tipo_sistema}', 0)
    data_atual = datetime.now().strftime("%d/%m/%Y")
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{organizacao},   {data_atual}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Detalhes do sistema
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"Detalhes do Sistema")
    font = run.font
    font.size = Pt(16)

    if endereco:
        document.add_paragraph(f"Localização: {endereco} {coordenadas}")
    else:
        document.add_paragraph(f"Coordenadas Geográficas: {coordenadas}")
    tipo_rede_curto = re.sub(r"\s*\([^)]*\)", "", tipo_rede)
    document.add_paragraph(f"Tipo da rede: {tipo_rede_curto}")
    document.add_paragraph(f"Orientação / Inclinação dos módulos: {orientacao} / {inclinacao}")
    seguranca = float(seguranca)*100
    document.add_paragraph(f"Fator de segurança (%) & coeficiente de temperatura (%/°C): {seguranca} & {temperatura}")
    document.add_paragraph(f"Autonomia máxima (h): {round(autonomia, 2)}")
    document.add_paragraph(f"Soma das potências (W): {int(potencia)}")
    document.add_paragraph(f"Consumo total (Wh): {int(consumo)}")

    document.add_paragraph(" ")
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"Dimensionamento Sugerido")
    font = run.font
    font.size = Pt(16)

    #document.add_paragraph(f"Eficiência do sistema (calculada): {eficiencia}")
    document.add_paragraph(f"Inversor recomendado: {dimensionamento.inversor}")
    document.add_paragraph(f"Quantidade de inversores: {dimensionamento.qtdinversor}")
    document.add_paragraph(f"Potência dos painéis: {painel}")
    document.add_paragraph(f"Quantidade de painéis: {dimensionamento.qtdplacas}")

    num_baterias = len(dimensionamento.baterias)
    if num_baterias == 1:
        bateria = list(dimensionamento.baterias.keys())[0]
        document.add_paragraph(f"Bateria Recomendada: {bateria}")
        document.add_paragraph(f"Quantidade de baterias: {dimensionamento.baterias[bateria]}")
    else:
        document.add_paragraph("Baterias Recomendadas:")
        for i, bateria in enumerate(dimensionamento.baterias):
            document.add_paragraph(f"Opção {i + 1}:")
            document.add_paragraph(f"   Modelo da bateria: {bateria}")
            document.add_paragraph(f"   Quantidade de baterias: {dimensionamento.baterias[bateria]}")

    if dimensionamento.autotrafo:
        document.add_paragraph(f"Necessário autotrafo: {dimensionamento.autotrafo}")

    # Adicionar uma quebra de página
    #if num_baterias == 1:
    document.add_page_break()

    tabela_geracao = dimensionamento.tabela_geracao

    if tabela_geracao is not None:
        consumo_anual = sum(tabela_consumo)
        geracao_anual = round(sum(tabela_geracao), 1)
        media_consumo = round(sum(tabela_consumo) / len(tabela_consumo), 1)
        media_geracao = round(sum(tabela_geracao) / len(tabela_geracao), 1)

        # Informações de consumo e geração
        document.add_paragraph(" ")
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Consumo e Geração")
        font = run.font
        font.size = Pt(16)
        document.add_paragraph(f"Média de consumo mensal (kWh): {media_consumo}")
        document.add_paragraph(f"Média de geração mensal (kWh): {media_geracao}")
        document.add_paragraph(f"Consumo anual de energia (kWh): {consumo_anual}")
        document.add_paragraph(f"Energia anual produzida (kWh): {geracao_anual}")
        document.add_paragraph(f"")

        # Dados fornecidos
        meses = np.arange(1, 13)  # Meses de 1 a 12
        plt.rcParams['font.family'] = 'Cambria'
        plt.rcParams.update({'font.size': 15})

        # Ajustar o tamanho do gráfico para ocupar toda a largura do documento
        plt.figure(figsize=(12, 5))  # Aumente a largura e altura conforme necessário

        # Configurações do gráfico
        bar_width = 0.35
        posicao_consumo = np.arange(len(meses))
        posicao_geracao = [x + bar_width for x in posicao_consumo]

        plt.bar(posicao_consumo, tabela_consumo, color='red', width=bar_width, label='Consumo')
        plt.bar(posicao_geracao, tabela_geracao, color='blue', width=bar_width, label='Geração')
        plt.xlabel('Mês')
        plt.ylabel('Quantidade (kWh)')
        plt.title('Consumo e Geração de Energia')
        plt.xticks([r + bar_width / 2 for r in range(len(meses))], [str(mes) for mes in meses])
        plt.legend()

        # Adicionando linhas pontilhadas horizontais
        valores_y = np.arange(0, max(max(tabela_consumo), max(tabela_geracao)) + 100,
                              100)  # Ajustando o espaçamento das linhas pontilhadas
        for valor in valores_y:
            plt.axhline(y=valor, color='gray', linestyle=(0, (5, 10)), linewidth=0.5)

        # Adicionando rodapé no gráfico
        plt.figtext(0.5, -0.1, "O valor da potência gerada pode variar de acordo com as condições climáticas de cada mês.",
                    ha="center", fontsize=15, bbox={"facecolor": "orange", "alpha": 0.5, "pad": 5})

        # Salvando o gráfico em um buffer temporário
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', bbox_inches="tight")
        buffer.seek(0)  # Voltar ao início do buffer

        # Adicionar o gráfico ao documento diretamente do buffer
        #document.add_paragraph("")
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(buffer, width=Inches(5.5))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        buffer.close()

        document.add_paragraph(f"")

        # Dividir os valores de HSP por 1000
        hsp = np.array(hsp) / 1000  # Convertendo para kWh

        # Criando o gráfico de linha para HSP
        plt.figure(figsize=(12, 5))  # Ajustando o tamanho do gráfico
        plt.plot(meses, hsp, marker='o', linestyle='-', color='blue', label='HSP')  # Linha com marcadores
        plt.xlabel('Mês')
        plt.ylabel('HSP (kWh/m²)')
        plt.title('Horas de Sol Pico no local (HSP)')
        plt.xticks(meses, [str(mes) for mes in meses])  # Meses no eixo X
        plt.yticks(np.arange(0, max(hsp) + 1, 1))  # Configurando os ticks de Y de 1 em 1
        plt.legend()

        # Adicionando linhas pontilhadas horizontais
        valores_y = np.arange(0, max(hsp) + 1, 1)  # Ajustando para valores de Y de 1 em 1
        for valor in valores_y:
            plt.axhline(y=valor, color='gray', linestyle=(0, (5, 10)), linewidth=0.5)

        # Adicionando rodapé
        plt.figtext(0.5, -0.1,
                    "Índices de irradiação de acordo com o Atlas Brasileiro de Energia Solar, INPE\nDisponível em: http://labren.ccst.inpe.br/",
                    ha="center", fontsize=15, bbox={"facecolor": "orange", "alpha": 0.5, "pad": 5})

        # Salvando o gráfico de HSP em um buffer temporário
        buffer_hsp = io.BytesIO()
        plt.savefig(buffer_hsp, format='png', bbox_inches="tight")
        buffer_hsp.seek(0)  # Voltar ao início do buffer

        # Adicionar o gráfico de HSP ao documento diretamente do buffer
        paragraph_hsp = document.add_paragraph()
        run_hsp = paragraph_hsp.add_run()
        run_hsp.add_picture(buffer_hsp, width=Inches(5.5))
        paragraph_hsp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Fechando o buffer ao finalizar
        buffer_hsp.close()

    # Adicionar parágrafo para o título
    document.add_paragraph(" ")
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Tabela de Cargas")
    run.font.size = Pt(16)

    # Criar uma tabela no documento Word
    num_cols = len(self.tree["columns"])
    num_rows = len(tabela)
    table = document.add_table(rows=num_rows + 1, cols=num_cols)

    # Preencher o cabeçalho da tabela
    header_cells = table.rows[0].cells
    for i, col_name in enumerate(self.tree["columns"]):
        if i == 2:
            col_name += " (W)"
        elif i == 3:
            col_name += " (h)"
        elif i == 4:
            col_name += " (Wh)"
        header_cells[i].text = col_name

    # Preencher as linhas da tabela com os dados copiados
    for row_index, row_data in enumerate(tabela):
        row_cells = table.rows[row_index + 1].cells
        for col_index, cell_data in enumerate(row_data):
            row_cells[col_index].text = str(cell_data)

    # Considerações finais
    document.add_paragraph(" ")
    document.add_paragraph(" ")
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"Considerações")
    font = run.font
    font.size = Pt(16)

    if tabela_geracao is not None:

        if modo_calculo == "Somar consumo + cargas":
            paragraph = document.add_paragraph(
                "A quantidade de módulos deste sistema foi calculada para atender tanto às necessidades de carga das baterias quanto ao consumo mensal do cliente. Isso resulta em uma quantidade maior de módulos, garantindo que ambas as demandas sejam satisfeitas, possivelmente proporcionando um excedente de energia.")
        else:
            paragraph = document.add_paragraph(
                "A quantidade de módulos deste sistema foi determinada pelo maior valor entre o consumo do cliente e a necessidade de carga das baterias. Isso resulta em uma quantidade menor de módulos, priorizando a maior demanda. No entanto, pode ser insuficiente para suprir completamente o consumo nos dias em que também é necessário recarregar as baterias.")

    paragraph2 = document.add_paragraph(
        "Este documento é de caráter informativo. A Vertys Solar Group não se responsabiliza pelas instalações elétricas. Cabe ao responsável técnico do projeto validar e aplicar as informações descritas de acordo com cada projeto, bem como avaliar a viabilidade de cada instalação.")

    # Definindo o alinhamento do parágrafo para justificado
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Salvar o arquivo
    nome_arquivo_padrao = "Dimensionamento"
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word files", "*.docx")],
        initialfile=nome_arquivo_padrao
    )
    if file_path:
        try:
            document.save(file_path)
            CustomInfoPopupWithOpen(self, "Sucesso", f"Relatório exportado para {file_path}", file_path)
        except PermissionError:
            popup = CustomInfoPopup(self, "Permissão negada",
                                    f"Por favor, feche o arquivo e tente novamente.")
            popup.grab_set()
